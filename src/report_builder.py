"""Build the X/Twitter report from results.json — as BOTH a .pdf and a .docx.

Only links that were captured cleanly are included; broken / failed links are
left out of the report entirely (how many were skipped is printed to console).

Layout:
  1. a HEADER at the top — the report title/date — with a separator rule under it
  2. the SCREENSHOTS — one clean tweet screenshot per link: category + handle +
     screenshot centered at the TOP of the page, link left-aligned beneath. The
     first tweet sits under the header; every later tweet starts its own page.
  3. a LINKS list — a single "Link" column (blue header), flowing right below
     the last screenshot block (no page break before it)

The title argument is the full header text (e.g. "Twitter Report 25-07-26").

Usage: python src/report_builder.py "Twitter Report 25-07-26"
"""
import html
import json
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "reports"


# --------------------------------------------------------------------------- #
# shared helpers
# --------------------------------------------------------------------------- #
def _has_categories(results):
    return any((r.get("category") or "Uncategorized") != "Uncategorized"
               for r in results)


def _shown_link(r):
    link = (r.get("post_link") or r.get("url") or "").strip()
    return "" if link.startswith("file://") else link


def _usable(r):
    """A link belongs in the report only if it was captured cleanly."""
    if r.get("status") != "ok":
        return False
    shot = r.get("screenshot")
    return bool(shot) and Path(shot).exists() and Path(shot).stat().st_size > 0


def _png_size(path):
    """(width, height) of a PNG from its IHDR — no image library needed."""
    with open(path, "rb") as f:
        head = f.read(24)
    if head[:8] == b"\x89PNG\r\n\x1a\n":
        import struct
        return struct.unpack(">II", head[16:24])
    return (0, 0)


# JPEG quality for the embedded screenshots. 88 with full chroma (subsampling=0)
# shrinks the file several-fold versus the raw PNGs while keeping tweet text
# crisp — no visible quality loss in the document.
_JPEG_QUALITY = 88


def _compress_for_embed(results, workdir):
    """Return a copy of `results` whose screenshots are high-quality JPEGs (much
    smaller to embed than PNGs), each tagged with its pixel size in `_dim`.
    Falls back to the original PNG if Pillow/convert fails."""
    try:
        from PIL import Image
    except Exception:
        return [{**r, "_dim": _png_size(r["screenshot"])} for r in results]

    out = []
    for i, r in enumerate(results, 1):
        src = r["screenshot"]
        try:
            im = Image.open(src).convert("RGB")
            dst = str(Path(workdir) / f"{i:03d}.jpg")
            im.save(dst, "JPEG", quality=_JPEG_QUALITY, optimize=True, subsampling=0)
            out.append({**r, "screenshot": dst, "_dim": im.size})
        except Exception:
            out.append({**r, "_dim": _png_size(src)})
    return out


def _fit(iw, ih, max_w, max_h):
    """Scale (iw, ih) to fit within (max_w, max_h) preserving aspect ratio."""
    if not iw or not ih:
        return max_w, max_w
    w = max_w
    h = w * ih / iw
    if h > max_h:
        h = max_h
        w = h * iw / ih
    return w, h


# =========================================================================== #
# PDF (reportlab)
# =========================================================================== #
def build_pdf(results, title, out):
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.utils import ImageReader
    from reportlab.platypus import (HRFlowable, Image, PageBreak, Paragraph,
                                    SimpleDocTemplate, Spacer, Table, TableStyle)

    HEADER_FILL = colors.HexColor("#D9E8F5")
    MARGIN = 0.75 * inch
    CONTENT_W = letter[0] - 2 * MARGIN
    SHOT_W = 4.9 * inch
    SHOT_MAX_H = 7.0 * inch

    ss = getSampleStyleSheet()
    HEADER = ParagraphStyle("XHeader", parent=ss["Title"], fontSize=18, leading=22,
                            alignment=TA_CENTER, spaceAfter=2)
    H1 = ParagraphStyle("XH1", parent=ss["Heading1"], fontSize=15)
    CATC = ParagraphStyle("XCat", parent=ss["Normal"], fontSize=11, alignment=TA_CENTER,
                          textColor=colors.HexColor("#666666"), spaceAfter=8)
    LINKL = ParagraphStyle("XLinkL", parent=ss["Normal"], fontSize=9, alignment=TA_LEFT,
                           spaceBefore=10, leading=12)
    CELL = ParagraphStyle("XCell", parent=ss["Normal"], fontSize=9, leading=12)

    def link_markup(url):
        u = html.escape(url, quote=True)
        return f'<a href="{u}" color="#1D9BF0"><u>{html.escape(url)}</u></a>'

    with_cat = _has_categories(results)
    flow = []

    # --- header + separator at the very top of the document
    flow.append(Paragraph(html.escape(title), HEADER))
    flow.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#999999"),
                           spaceBefore=4, spaceAfter=14))

    # --- screenshots: category + handle + shot centered at the TOP, link
    #     left-aligned beneath. The first tweet sits under the header; every
    #     later tweet starts its own page.
    for j, r in enumerate(results):
        if j:
            flow.append(PageBreak())
        if with_cat:
            flow.append(Paragraph(html.escape(r.get("category") or "Uncategorized"), CATC))

        iw, ih = r.get("_dim") or ImageReader(r["screenshot"]).getSize()
        w, h = _fit(iw, ih, min(SHOT_W, CONTENT_W), SHOT_MAX_H)
        img = Image(r["screenshot"], width=w, height=h)
        img.hAlign = "CENTER"
        flow.append(img)

        link = _shown_link(r)
        if link:
            flow.append(Paragraph("Link — " + link_markup(link), LINKL))

    # --- links list: single "Link" column (blue header), flowing right below the
    #     last screenshot block (no page break before it).
    flow += [Spacer(1, 18), Paragraph("Links", H1), Spacer(1, 8)]
    data = [[Paragraph('<b><font color="#1D9BF0">Link</font></b>', CELL)]]
    for r in results:
        link = _shown_link(r)
        data.append([Paragraph(link_markup(link), CELL) if link else Paragraph("—", CELL)])
    tbl = Table(data, colWidths=[CONTENT_W], hAlign="CENTER")
    tbl.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#BBBBBB")),
        ("BACKGROUND", (0, 0), (-1, 0), HEADER_FILL),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    flow.append(tbl)

    doc = SimpleDocTemplate(str(out), pagesize=letter, leftMargin=MARGIN,
                            rightMargin=MARGIN, topMargin=MARGIN, bottomMargin=MARGIN,
                            title=title)
    doc.build(flow)


# =========================================================================== #
# DOCX (python-docx)
# =========================================================================== #
def build_docx(results, title, out):
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Inches, Pt, RGBColor

    CENTER = WD_ALIGN_PARAGRAPH.CENTER

    def set_valign(section, val):
        sectPr = section._sectPr
        for e in sectPr.findall(qn("w:vAlign")):
            sectPr.remove(e)
        v = OxmlElement("w:vAlign"); v.set(qn("w:val"), val); sectPr.append(v)

    def shade(cell, fill):
        tcpr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), fill)
        tcpr.append(shd)

    def add_rule(paragraph):
        """Draw a horizontal separator line beneath a paragraph."""
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "6"); bottom.set(qn("w:color"), "999999")
        pBdr.append(bottom); pPr.append(pBdr)

    def hyperlink(paragraph, url, text):
        part = paragraph.part
        r_id = part.relate_to(
            url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True)
        h = OxmlElement("w:hyperlink"); h.set(qn("r:id"), r_id)
        run = OxmlElement("w:r"); rpr = OxmlElement("w:rPr")
        color = OxmlElement("w:color"); color.set(qn("w:val"), "1D9BF0"); rpr.append(color)
        u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rpr.append(u)
        run.append(rpr)
        t = OxmlElement("w:t"); t.text = text; run.append(t)
        h.append(run); paragraph._p.append(h)

    with_cat = _has_categories(results)
    doc = Document()

    # --- header + separator at the very top of the document
    set_valign(doc.sections[0], "top")
    hp = doc.add_paragraph(); hp.alignment = CENTER
    hrun = hp.add_run(title); hrun.bold = True; hrun.font.size = Pt(18)
    add_rule(hp)

    # --- screenshots: category + handle + shot centered at the TOP, link
    #     left-aligned beneath. The first tweet sits under the header; every
    #     later tweet starts its own page.
    for j, r in enumerate(results):
        if j:
            sec = doc.add_section(WD_SECTION.NEW_PAGE)
            set_valign(sec, "top")
        if with_cat:
            cp = doc.add_paragraph(); cp.alignment = CENTER
            cr = cp.add_run(r.get("category") or "Uncategorized")
            cr.italic = True; cr.font.size = Pt(10); cr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        pp = doc.add_paragraph(); pp.alignment = CENTER
        iw, ih = r.get("_dim") or _png_size(r["screenshot"])
        w_in, h_in = _fit(iw, ih, 4.9, 7.0)
        pp.add_run().add_picture(r["screenshot"], width=Inches(w_in), height=Inches(h_in))

        link = _shown_link(r)
        if link:
            lp = doc.add_paragraph(); lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            lp.add_run("Link — ").font.size = Pt(9)
            hyperlink(lp, link, link)

    # --- links list: single "Link" column (blue header), flowing right below the
    #     last screenshot block (no new section / no page break).
    doc.add_heading("Links", level=1)
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells[0]
    hrun = hdr.paragraphs[0].add_run("Link")                            # fresh cell, no empty run
    hrun.bold = True; hrun.font.color.rgb = RGBColor(0x1D, 0x9B, 0xF0)   # blue header text
    shade(hdr, "D9E8F5")                                                # fill unchanged
    for r in results:
        cell = table.add_row().cells[0]
        link = _shown_link(r); cell.text = ""
        if link:
            hyperlink(cell.paragraphs[0], link, link)
        else:
            cell.text = "—"

    doc.save(out)


def main():
    import tempfile

    title = sys.argv[1] if len(sys.argv) > 1 else "X Report"
    stem = sys.argv[2] if len(sys.argv) > 2 else "X_Report"
    all_results = json.loads((OUT / "results.json").read_text())
    results = [r for r in all_results if _usable(r)]
    skipped = len(all_results) - len(results)

    if not results:
        print(f"[report] no capturable links ({skipped} skipped) — nothing to build")
        return

    pdf = OUT / f"{stem}.pdf"
    docx = OUT / f"{stem}.docx"
    with tempfile.TemporaryDirectory() as td:
        embed = _compress_for_embed(results, td)   # shrink the embedded images
        build_pdf(embed, title, pdf)
        build_docx(embed, title, docx)
    extra = f"  ({skipped} skipped)" if skipped else ""
    print(f"[report] wrote {pdf}  ({_mb(pdf)} MB)")
    print(f"[report] wrote {docx}  ({_mb(docx)} MB, {len(results)} link(s){extra})")


def _mb(path):
    return round(Path(path).stat().st_size / 1_048_576, 1)


if __name__ == "__main__":
    main()
