"""Build the Influencer report from results.json — as BOTH a .pdf and a .docx,
with the same layout in each. A4 pages, TWO posts per page side by side.

`src/report_builder.py` is frozen and untouched; this is its parallel.

Layout:
  1. HEADER — the report title, centered, with an accent rule beneath it.
  2. A two-column grid, two posts per page, in input order. Each post is:
       * the ACCOUNT NAME as a coloured heading;
       * the SCREENSHOT (username + text + media + likes/reposts) in a bordered
         card;
       * METRICS as label -> value rows, hairline-ruled rather than boxed:
         Followers, Reactions, Comments, Reach, Shares;
       * a clickable "Link: <url>" line.
  3. LINKS TABLE at the very end: every link in input order, single column,
     NO serial numbers, each entry clickable.

Only cleanly captured posts are included; failures are dropped (and reported by
the web UI's activity log).

Usage: python influencer/inf_report_builder.py "Influencer Report 25-07-26" stem
"""
import html
import json
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "reports"

# --------------------------------------------------------------------------- #
# Design tokens — one accent on a neutral slate palette
# --------------------------------------------------------------------------- #
ACCENT = "#1D9BF0"      # links + the header rule
HEADING = "#B23A21"     # the account name above each post
INK = "#0F172A"
MUTED = "#64748B"
RULE = "#E2E8F0"
TINT = "#F8FAFC"

# Audience first, then engagement. Reach = Views, Shares = Reposts.
METRIC_FIELDS = (("Followers", "followers"), ("Reactions", "reactions"),
                 ("Comments", "comments"), ("Reach", "reach"),
                 ("Shares", "shares"))
MISSING = "—"

_JPEG_QUALITY = 88

POSTS_PER_PAGE = 2

# Page geometry, shared by the PDF and the DOCX so both lay out identically.
# A 0.6 in margin buys noticeably larger screenshots than 0.7 in once the page
# is split into two columns.
MARGIN_IN = 0.6
COL_IN = 3.3            # usable width inside one column, after the gutter

# Max screenshot size in inches (width, height) inside one column. The height cap
# leaves room for the heading, the five metric rows and the link line, so a pair
# always fits on one page.
SHOT_MAX_IN = (3.05, 6.9)

# A Unicode-capable TTF, so a non-Latin report title (e.g. Hindi) renders in the
# PDF instead of coming out as black boxes. Falls back to Helvetica.
_FONT_CANDIDATES = [
    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
    "/usr/share/fonts/TTF/DejaVuSans.ttf",
    "/Library/Fonts/Arial Unicode.ttf",
    "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
]


# --------------------------------------------------------------------------- #
# Shared helpers
# --------------------------------------------------------------------------- #
def _usable(r) -> bool:
    """A post belongs in the report only if it was captured cleanly."""
    if r.get("status") != "ok":
        return False
    shot = r.get("screenshot")
    return bool(shot) and Path(shot).exists() and Path(shot).stat().st_size > 0


def _shown_link(r) -> str:
    link = (r.get("post_link") or r.get("url") or "").strip()
    return "" if link.startswith("file://") else link


def _metric(r, key) -> str:
    value = ((r.get("metrics") or {}).get(key) or "").strip()
    return value or MISSING


def _account_label(r) -> str:
    """The heading shown above a post — the account, never the input sheet's
    category (which is often just a section word like "Tweet links").

    Prefers a real name from the input sheet, falls back to the @handle the
    capture read off the page.
    """
    name = (r.get("account_name") or "").strip()
    handle = (r.get("handle") or "").strip()
    placeholder = name.lower() in ("", "x post") or name.startswith("@")
    if not placeholder:
        return name
    return handle or name


def _pairs(results):
    """results -> [(a, b), (c, d), …]; the last pair may have a single post."""
    return [results[i:i + POSTS_PER_PAGE]
            for i in range(0, len(results), POSTS_PER_PAGE)]


def _png_size(path):
    """(width, height) from a PNG's IHDR — no image library needed."""
    try:
        with open(path, "rb") as f:
            head = f.read(24)
        if head[:8] == b"\x89PNG\r\n\x1a\n":
            import struct
            return struct.unpack(">II", head[16:24])
    except OSError:
        pass
    return (0, 0)


def _compress_for_embed(results, workdir):
    """High-quality JPEGs are several times smaller to embed than the raw PNGs
    with no visible loss. Falls back to the PNG if Pillow is unavailable."""
    try:
        from PIL import Image
    except Exception:
        return [{**r, "_dim": _png_size(r["screenshot"])} for r in results]

    out = []
    for i, r in enumerate(results, 1):
        src = r["screenshot"]
        try:
            im = Image.open(src).convert("RGB")
            dst = str(Path(workdir) / f"{i:03d}.jpg")
            im.save(dst, "JPEG", quality=_JPEG_QUALITY, optimize=True, subsampling=0)
            out.append({**r, "screenshot": dst, "_dim": im.size})
        except Exception:
            out.append({**r, "_dim": _png_size(src)})
    return out


def _fit(iw, ih, max_w, max_h):
    """Scale (iw, ih) into (max_w, max_h), preserving aspect ratio."""
    if not iw or not ih:
        return max_w, max_w
    w = max_w
    h = w * ih / iw
    if h > max_h:
        h = max_h
        w = h * iw / ih
    return w, h


# =========================================================================== #
# PDF (reportlab)
# =========================================================================== #
def _register_unicode_font():
    """Return (regular, bold) font names, preferring a Unicode TTF."""
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    for path in _FONT_CANDIDATES:
        p = Path(path)
        if not p.exists():
            continue
        try:
            pdfmetrics.registerFont(TTFont("ReportSans", str(p)))
            bold = p.with_name(p.name.replace("DejaVuSans.ttf",
                                              "DejaVuSans-Bold.ttf"))
            if bold != p and bold.exists():
                pdfmetrics.registerFont(TTFont("ReportSans-Bold", str(bold)))
                return "ReportSans", "ReportSans-Bold"
            return "ReportSans", "ReportSans"
        except Exception:
            continue
    return "Helvetica", "Helvetica-Bold"


def build_pdf(results, title, out):
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.utils import ImageReader
    from reportlab.platypus import (HRFlowable, Image, PageBreak, Paragraph,
                                    SimpleDocTemplate, Spacer, Table, TableStyle)

    FONT, BOLD = _register_unicode_font()

    MARGIN = MARGIN_IN * inch
    CONTENT_W = A4[0] - 2 * MARGIN
    GUTTER = 9                       # half-gutter of padding on each column
    COL_W = CONTENT_W / 2.0
    INNER_W = COL_W - 2 * GUTTER     # usable width inside one column
    CARD_PAD = 7
    SHOT_W = min(SHOT_MAX_IN[0] * inch, INNER_W - 2 * CARD_PAD)
    SHOT_MAX_H = SHOT_MAX_IN[1] * inch

    HEADER = ParagraphStyle("Header", fontName=BOLD, fontSize=19, leading=24,
                            alignment=TA_CENTER, textColor=colors.HexColor(INK),
                            spaceAfter=2)
    ACCOUNT = ParagraphStyle("Account", fontName=BOLD, fontSize=11.5, leading=15,
                             alignment=TA_LEFT,
                             textColor=colors.HexColor(HEADING), spaceAfter=6)
    MLABEL = ParagraphStyle("MLabel", fontName=BOLD, fontSize=10.5, leading=14,
                            alignment=TA_LEFT, textColor=colors.HexColor(MUTED))
    LINKL = ParagraphStyle("LinkLine", fontName=FONT, fontSize=7.6, leading=10.5,
                           alignment=TA_LEFT, textColor=colors.HexColor(MUTED),
                           spaceBefore=9)
    CELL = ParagraphStyle("Cell", fontName=FONT, fontSize=9, leading=13,
                          textColor=colors.HexColor(INK))
    H1 = ParagraphStyle("H1", fontName=BOLD, fontSize=13, leading=17,
                        textColor=colors.HexColor(INK), spaceAfter=8)

    def link_markup(url, text=None):
        safe = html.escape(url, quote=True)
        return (f'<a href="{safe}" color="{ACCENT}">'
                f'<u>{html.escape(text or url)}</u></a>')

    def post_block(r):
        """One post as a list of flowables, sized to a single column."""
        block = []

        label = _account_label(r)
        if label:
            block.append(Paragraph(html.escape(label), ACCOUNT))

        iw, ih = r.get("_dim") or ImageReader(r["screenshot"]).getSize()
        w, h = _fit(iw, ih, SHOT_W, SHOT_MAX_H)
        img = Image(r["screenshot"], width=w, height=h)
        img.hAlign = "CENTER"
        card = Table([[img]], colWidths=[INNER_W], hAlign="LEFT")
        card.setStyle(TableStyle([
            ("BOX", (0, 0), (-1, -1), 0.75, colors.HexColor(RULE)),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("TOPPADDING", (0, 0), (-1, -1), CARD_PAD),
            ("BOTTOMPADDING", (0, 0), (-1, -1), CARD_PAD),
            ("LEFTPADDING", (0, 0), (-1, -1), CARD_PAD),
            ("RIGHTPADDING", (0, 0), (-1, -1), CARD_PAD),
        ]))
        block += [card, Spacer(1, 11)]

        # metrics: inline "LABEL  value" rows, hairline-ruled instead of boxed.
        # Inline (rather than a right-aligned second column) so the DOCX can be
        # built from plain paragraphs and match this exactly — see build_docx.
        rows = [[Paragraph(
            f'<font size="7.5" color="{MUTED}">{name.upper()}</font>'
            f'<font size="10.5" color="{INK}">&nbsp;&nbsp;'
            f'{html.escape(_metric(r, key))}</font>', MLABEL)]
            for name, key in METRIC_FIELDS]
        metrics = Table(rows, colWidths=[INNER_W], hAlign="LEFT")
        metrics.setStyle(TableStyle([
            ("LINEBELOW", (0, 0), (-1, -2), 0.4, colors.HexColor(RULE)),
            ("LINEABOVE", (0, 0), (-1, 0), 1.0, colors.HexColor(INK)),
            # BOTTOM, not MIDDLE: the label and value are different sizes, and
            # middle-aligning two different text heights leaves the value
            # visibly sitting below its label.
            ("VALIGN", (0, 0), (-1, -1), "BOTTOM"),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("LEFTPADDING", (0, 0), (-1, -1), 0),
            ("RIGHTPADDING", (0, 0), (-1, -1), 0),
        ]))
        block.append(metrics)

        link = _shown_link(r)
        if link:
            block.append(Paragraph(
                f'<font color="{INK}"><b>Link:</b></font> ' + link_markup(link),
                LINKL))
        return block

    flow = []

    # 1) Header ------------------------------------------------------------
    flow.append(Paragraph(html.escape(title), HEADER))
    flow.append(HRFlowable(width="100%", thickness=1.6,
                           color=colors.HexColor(ACCENT),
                           spaceBefore=6, spaceAfter=3))
    flow.append(HRFlowable(width="100%", thickness=0.6,
                           color=colors.HexColor(RULE),
                           spaceBefore=0, spaceAfter=16))

    # 2+3) Two posts per page, side by side --------------------------------
    for n, pair in enumerate(_pairs(results)):
        if n:
            flow.append(PageBreak())
        cells = [post_block(r) for r in pair]
        while len(cells) < POSTS_PER_PAGE:
            cells.append("")                     # a lone final post
        grid = Table([cells], colWidths=[COL_W] * POSTS_PER_PAGE, hAlign="CENTER")
        grid.setStyle(TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), GUTTER),
            ("RIGHTPADDING", (0, 0), (-1, -1), GUTTER),
            ("TOPPADDING", (0, 0), (-1, -1), 0),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
        ]))
        flow.append(grid)

    # 4) Links table -------------------------------------------------------
    flow += [PageBreak(), Paragraph("Links", H1)]
    data = [[Paragraph(f'<font color="{ACCENT}"><b>Link</b></font>', CELL)]]
    for r in results:
        link = _shown_link(r)
        data.append([Paragraph(link_markup(link) if link else MISSING, CELL)])
    table = Table(data, colWidths=[CONTENT_W], hAlign="CENTER", repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(TINT)),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor(RULE)),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
    ]))
    flow.append(table)

    doc = SimpleDocTemplate(str(out), pagesize=A4, leftMargin=MARGIN,
                            rightMargin=MARGIN, topMargin=MARGIN,
                            bottomMargin=MARGIN, title=title)
    doc.build(flow)


# =========================================================================== #
# DOCX (python-docx) — the same layout, A4
# =========================================================================== #
def build_docx(results, title, out):
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Mm, Pt, RGBColor

    CENTER = WD_ALIGN_PARAGRAPH.CENTER

    def rgb(hex_str):
        h = hex_str.lstrip("#")
        return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

    def a4(section):
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        for side in ("left", "right", "top", "bottom"):
            setattr(section, f"{side}_margin", Inches(MARGIN_IN))
        sectPr = section._sectPr
        for e in sectPr.findall(qn("w:vAlign")):
            sectPr.remove(e)
        v = OxmlElement("w:vAlign")
        v.set(qn("w:val"), "top")
        sectPr.append(v)

    def shade(cell, fill):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), fill.lstrip("#"))
        tcPr.append(shd)

    def rule(paragraph, color, size="12"):
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), size)
        bottom.set(qn("w:space"), "6")
        bottom.set(qn("w:color"), color.lstrip("#"))
        pBdr.append(bottom)
        pPr.append(pBdr)

    def hyperlink(paragraph, url, text, size=9):
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/"
            "relationships/hyperlink",
            is_external=True)
        link = OxmlElement("w:hyperlink")
        link.set(qn("r:id"), r_id)
        run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), ACCENT.lstrip("#"))
        rPr.append(color)
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(size * 2)))
        rPr.append(sz)
        run.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        run.append(t)
        link.append(run)
        paragraph._p.append(link)

    def para_border(paragraph, edges, color, size="4", space="4"):
        """Draw border edges on a paragraph (python-docx has no API for this).

        Used instead of nested tables: paragraph borders and tab stops are laid
        out identically by every Word-compatible renderer, whereas a nested
        table's width depends on the renderer honouring w:tblGrid.
        """
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = pPr.find(qn("w:pBdr"))
        if pBdr is None:
            pBdr = OxmlElement("w:pBdr")
            pPr.append(pBdr)
        for edge in edges:
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), size)
            el.set(qn("w:space"), space)
            el.set(qn("w:color"), color.lstrip("#"))
            pBdr.append(el)

    def tight(paragraph, before=0, after=0):
        paragraph.paragraph_format.space_before = Pt(before)
        paragraph.paragraph_format.space_after = Pt(after)
        return paragraph

    def fixed_width(table, col_widths_in):
        """Pin a table to exact column widths.

        Three things are needed, and missing any one of them leaves the table
        sized to its text instead of spanning the page:
          * a fixed layout plus an explicit total width (w:tblW);
          * per-cell widths (w:tcW);
          * the table GRID (w:tblGrid/w:gridCol) — under a fixed layout Word
            takes its column widths from the grid, and python-docx never
            updates it when you set `cell.width`.
        """
        table.autofit = False
        twips = [int(w * 1440) for w in col_widths_in]

        tblPr = table._tbl.tblPr
        for tag in ("w:tblW", "w:tblLayout"):
            existing = tblPr.find(qn(tag))
            if existing is not None:
                tblPr.remove(existing)
        tblW = OxmlElement("w:tblW")
        tblW.set(qn("w:type"), "dxa")
        tblW.set(qn("w:w"), str(sum(twips)))
        tblPr.append(tblW)
        layout = OxmlElement("w:tblLayout")
        layout.set(qn("w:type"), "fixed")
        tblPr.append(layout)

        grid = table._tbl.find(qn("w:tblGrid"))
        if grid is not None:
            for col, width in zip(grid.findall(qn("w:gridCol")), twips):
                col.set(qn("w:w"), str(width))

        for row in table.rows:
            for cell, width in zip(row.cells, col_widths_in):
                cell.width = Inches(width)

    def post_block(cell, r):
        """Render one post into an outer-grid cell — mirrors post_block() in
        the PDF so the two documents look the same."""
        # Everything below is paragraphs only — no nested tables — so the block
        # lays out the same in Word, WPS and Google Docs.
        label = _account_label(r)
        first = tight(cell.paragraphs[0], after=5)
        if label:
            arun = first.add_run(label)
            arun.bold = True
            arun.font.size = Pt(11.5)
            arun.font.color.rgb = rgb(HEADING)

        # screenshot inside a bordered card (a boxed paragraph)
        spar = tight(cell.add_paragraph(), after=9)
        spar.alignment = CENTER
        para_border(spar, ("top", "bottom", "left", "right"), RULE, "6", space="5")
        iw, ih = r.get("_dim") or _png_size(r["screenshot"])
        w_in, h_in = _fit(iw, ih, *SHOT_MAX_IN)
        spar.add_run().add_picture(r["screenshot"], width=Inches(w_in),
                                   height=Inches(h_in))

        # metrics: inline "LABEL  value" per paragraph with a hairline rule
        # underneath — same as the PDF. Two runs in one left-aligned paragraph,
        # so there is no tab stop or nested-table width for a renderer to
        # interpret differently.
        last = len(METRIC_FIELDS) - 1
        for i, (name, key) in enumerate(METRIC_FIELDS):
            mp = tight(cell.add_paragraph(), before=2, after=2)
            if i == 0:
                para_border(mp, ("top",), INK, "10", space="2")
            if i != last:
                para_border(mp, ("bottom",), RULE, "4", space="2")

            lrun = mp.add_run(name.upper())
            lrun.bold = True
            lrun.font.size = Pt(7.5)
            lrun.font.color.rgb = rgb(MUTED)

            vrun = mp.add_run("  " + _metric(r, key))
            vrun.bold = True
            vrun.font.size = Pt(10.5)
            vrun.font.color.rgb = rgb(INK)

        link = _shown_link(r)
        if link:
            lp = tight(cell.add_paragraph(), before=7)
            lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            lrun = lp.add_run("Link: ")
            lrun.bold = True
            lrun.font.size = Pt(7.6)
            lrun.font.color.rgb = rgb(INK)
            hyperlink(lp, link, link, size=7.6)

    doc = Document()
    a4(doc.sections[0])

    # 1) Header ------------------------------------------------------------
    hp = doc.add_paragraph()
    hp.alignment = CENTER
    hrun = hp.add_run(title)
    hrun.bold = True
    hrun.font.size = Pt(19)
    hrun.font.color.rgb = rgb(INK)
    rule(hp, ACCENT, "16")

    # 2+3) Two posts per page, side by side --------------------------------
    pairs = _pairs(results)
    for n, pair in enumerate(pairs):
        if n:
            doc.add_page_break()
        grid = doc.add_table(rows=1, cols=POSTS_PER_PAGE)
        grid.alignment = WD_TABLE_ALIGNMENT.CENTER
        fixed_width(grid, [COL_IN + 0.12] * POSTS_PER_PAGE)
        for col, r in enumerate(pair):
            post_block(grid.rows[0].cells[col], r)

    # 4) Links table -------------------------------------------------------
    # Own page: with two posts per page the last grid already fills the page,
    # so letting the table spill would split it awkwardly.
    doc.add_page_break()
    lh = doc.add_paragraph()
    lrun = lh.add_run("Links")
    lrun.bold = True
    lrun.font.size = Pt(13)
    lrun.font.color.rgb = rgb(INK)

    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    head = table.rows[0].cells[0]
    hrun3 = head.paragraphs[0].add_run("Link")
    hrun3.bold = True
    hrun3.font.color.rgb = rgb(ACCENT)
    shade(head, TINT)
    for r in results:
        cell = table.add_row().cells[0]
        link = _shown_link(r)
        cell.text = ""
        if link:
            hyperlink(cell.paragraphs[0], link, link)
        else:
            cell.text = MISSING

    doc.save(out)


# =========================================================================== #
def _mb(path):
    return round(Path(path).stat().st_size / 1_048_576, 1)


def main():
    import tempfile

    title = sys.argv[1] if len(sys.argv) > 1 else "Influencer Report"
    stem = sys.argv[2] if len(sys.argv) > 2 else "Influencer_Report"

    all_results = json.loads((OUT / "results.json").read_text())
    results = [r for r in all_results if _usable(r)]
    skipped = len(all_results) - len(results)

    if not results:
        print(f"[report] no capturable links ({skipped} skipped) — nothing to build")
        return

    pdf = OUT / f"{stem}.pdf"
    docx = OUT / f"{stem}.docx"
    with tempfile.TemporaryDirectory() as td:
        embed = _compress_for_embed(results, td)
        build_pdf(embed, title, pdf)
        build_docx(embed, title, docx)

    extra = f"  ({skipped} skipped)" if skipped else ""
    print(f"[report] wrote {pdf}  ({_mb(pdf)} MB)")
    print(f"[report] wrote {docx}  ({_mb(docx)} MB, {len(results)} link(s){extra})")


if __name__ == "__main__":
    main()
